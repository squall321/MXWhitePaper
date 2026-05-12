"""DocumentJSON v1.0 → Word (.docx) 출력.

`POST /exports/docx` 가 호출하는 단방향 익스포터. cycle 1 의 docx_import 와
짝을 이루어 round-trip 가능 (importer 가 인식 가능한 스타일/구조를 유지).

설계 요약
---------
- python-docx 사용 (pyproject 에 이미 포함). 다른 export(`pptx_export`,
  `markdown_export`) 와 동일한 doc-walk 패턴.
- title slide → Title style + Subtitle (summary). metadata 는 italic 한 줄.
- Heading 1/2/3 = section level, Heading 4/5/6 = `heading-4` block.
- markdown-lite (`**bold**` / `*italic*` / `~~strike~~` / `__underline__` /
  `` `code` `` / `[label](url)` / `[[slug|label]]` / `[^N]`) 는
  ``_emit_inline_runs`` 가 한 번에 토큰화해 docx run/hyperlink/footnote 로 변환.
- list block → 'List Bullet' / 'List Number' 스타일 + 들여쓰기로 depth 표현.
- table block → 실제 ``Table`` 도형, header row bold. ``meta.note`` caption →
  표 직후 Caption 스타일 단락.
- callout → 단락 shading XML 으로 variant color 표시 + 좌측 boundary.
- code → Consolas + 회색 배경, language 가 있으면 위에 작은 caption.
- math → ``$expr$`` 폴백 (OMML 임베드는 v1 범위 밖, 라이브러리 한계 참조).
- image → ``image_resolver`` 로 bytes 가져와 ``add_picture``. 실패시 텍스트 폴백.
- chart / gantt / flow / org-chart → 텍스트 요약 + placeholder note.
- iframe / video / file / dashboard-embed / data-source / calculator →
  제목 + URL 단락.
- tabs / accordion / columns → 평탄화 (heading + 본문 sequential).
- ``meta.note: 'speaker:…'`` → DROP (DOCX 노트 페인 없음).
- ``meta.note: 'page-break-before'`` → ``add_page_break()``.

알려진 한계
-----------
- math 블록은 OMML 미지원 (python-docx 가 ``<m:oMath>`` 직접 빌드 API 를
  제공하지 않음). 호출자가 ``$expr$`` 텍스트 폴백을 인식해야 한다.
- chart / gantt / flow / org-chart 는 시각적 요소가 텍스트 outline 으로
  대체된다 (round-trip 시 text 로 보존되지만 도형은 손실).
- callout 색상은 paragraph shading 만 지원. 좌측 보더는 paragraph border
  XML 으로 직접 구성.
- footnote 정의(`[^N]: …`) 는 본문 어디에 있든 전역 수집되어 add_footnote()
  로 변환. 정의 자체는 별도 단락으로 출력하지 않는다.

부수효과 없음 — 입력 dict 변경하지 않으며, 결과는 BytesIO 의 .docx 바이트.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Tunables ─────────────────────────────────────────────────────────

# Callout variant → background hex (light tints) + accent border color.
_CALLOUT_COLORS: dict[str, tuple[str, str]] = {
    "info": ("DBE4F0", "1F3AA8"),    # Samsung blue
    "warn": ("FFF7E0", "B45309"),    # amber
    "danger": ("FCE2E0", "B91C1C"),  # red
    "tip": ("DDF7E8", "047857"),     # green
}

# `code` block background (light gray).
_CODE_SHADING = "F1F5F9"

# Inline `code` token shading.
_INLINE_CODE_SHADING = "EEEEEE"

# Markdown-lite token regex — order matters (longer ops first).
_INLINE_RE = re.compile(
    r"(?P<bold>\*\*([^*]+)\*\*)"
    r"|(?P<strike>~~([^~]+)~~)"
    r"|(?P<underline>__([^_]+)__)"
    r"|(?P<italic>(?<![*])\*([^*]+)\*(?!\*))"
    r"|(?P<code>`([^`]+)`)"
    r"|(?P<link>\[([^\]]+)\]\(([^)]+)\))"
    r"|(?P<wiki>\[\[([^\]]+)\]\])"
    r"|(?P<footnote>\[\^([^\]]+)\])"
)

ImageResolver = Callable[[str], dict[str, Any] | None]


# ── Public entry ─────────────────────────────────────────────────────


@dataclass
class DocxOptions:
    """렌더 옵션.

    ``image_resolver``: image_id → ``{bytes: bytes, mime: str}`` 또는 None.
    """

    image_resolver: ImageResolver | None = None


def render_docx(
    doc: dict[str, Any],
    *,
    options: DocxOptions | None = None,
    requester_role: str | None = None,
) -> bytes:
    """DocumentJSON dict 를 .docx 바이너리로 렌더한다.

    Args:
        doc: DocumentJSON v1.0 dict.
        options: 렌더 옵션. None 이면 기본값.
        requester_role: 호출자 role. 지정 시 admin 미만은 meta.permission 이
            높은 블록을 redact 한 결과를 렌더한다. None 이면 scrub 미적용.

    Returns:
        .docx zip 바이너리 (PK\\x03\\x04 매직).
    """
    opts = options or DocxOptions()
    if requester_role is not None:
        from .document_service import scrub_for_response

        doc = scrub_for_response(doc, role=requester_role)
    document = Document()
    ctx = _Ctx(document=document, opts=opts, footnotes=_collect_footnotes(doc), requester_role=requester_role)

    _render_title(document, doc)
    for section in doc.get("sections") or []:
        _render_section(document, section, ctx)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ── Internal context ────────────────────────────────────────────────


@dataclass
class _Ctx:
    document: Any
    opts: DocxOptions
    footnotes: dict[str, str]  # marker -> body text
    figure_counter: int = 0     # incremented each time an image with caption is rendered
    requester_role: str | None = None


# ── Title / metadata ────────────────────────────────────────────────


def _render_title(document: Any, doc: dict[str, Any]) -> None:
    title = _str(doc.get("title")) or "Untitled"
    summary = _str(doc.get("summary"))
    metadata = doc.get("metadata") or {}

    p = document.add_paragraph(title, style="Title")
    if summary:
        document.add_paragraph(summary, style="Subtitle")

    meta_parts: list[str] = []
    for key in ("division", "team", "group", "part"):
        v = _str(metadata.get(key))
        if v:
            meta_parts.append(v)
    owners = metadata.get("owners") or []
    if owners:
        meta_parts.append("owners: " + ", ".join(_str(o) for o in owners))
    tags = metadata.get("tags") or []
    if tags:
        meta_parts.append("tags: " + ", ".join(_str(t) for t in tags))
    if meta_parts:
        meta_p = document.add_paragraph()
        run = meta_p.add_run(" / ".join(meta_parts))
        run.italic = True
        run.font.size = Pt(10)


# ── Section rendering ────────────────────────────────────────────────


def _render_section(document: Any, section: dict[str, Any], ctx: _Ctx) -> None:
    level = int(section.get("level") or 1)
    number = _str(section.get("number"))
    title = _str(section.get("title"))
    heading_text = f"{number} {title}".strip()
    heading_level = max(1, min(level, 3))
    document.add_heading(heading_text, level=heading_level)

    for block in section.get("blocks") or []:
        _render_block(document, block, ctx)
    for sub in section.get("subsections") or []:
        _render_section(document, sub, ctx)


# ── Block dispatch ───────────────────────────────────────────────────


def _render_block(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    btype = _str(block.get("type"))
    meta = block.get("meta") if isinstance(block.get("meta"), dict) else {}
    note = _str(meta.get("note")) if meta else ""

    # `slide-only` blocks belong on a presentation deck, not in a Word
    # document; drop them so the .docx export reads naturally even when
    # the author keeps presenter-visual blocks in the document body.
    audience = _str(meta.get("audience")) if meta else ""
    if audience == "slide-only":
        return

    # speaker notes have no DOCX equivalent — drop.
    if note.startswith("speaker:"):
        return

    # page-break-before — add_page_break BEFORE rendering.
    if note == "page-break-before":
        # Empty paragraph carrier — emit page break and return.
        if btype == "paragraph" and not _str(block.get("text")):
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            return
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        # fallthrough — render the block normally after the break.

    # Footnote definition lines `[^N]: …` are hoisted via add_footnote — skip
    # the literal definition paragraph in the output.
    if btype == "paragraph":
        text = _str(block.get("text"))
        if _is_footnote_def(text):
            return

    handler = _BLOCK_HANDLERS.get(btype)
    if handler is None:
        p = document.add_paragraph()
        run = p.add_run(f"[알 수 없는 블록: {btype}]")
        run.italic = True
        return
    try:
        handler(document, block, ctx)
    except Exception as e:  # pragma: no cover — defensive
        p = document.add_paragraph()
        run = p.add_run(f"[블록 렌더 실패: {btype} — {e}]")
        run.italic = True


# ── Per-block handlers ───────────────────────────────────────────────


def _b_paragraph(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    text = _str(block.get("text"))
    if not text:
        return
    p = document.add_paragraph()
    _emit_inline_runs(p, text, ctx)


def _b_heading_4(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    title = _str(block.get("title"))
    if not title:
        return
    meta = block.get("meta") or {}
    level = 4
    if isinstance(meta, dict):
        ml = meta.get("level")
        if ml in (2, 3, 4):
            level = int(ml) + 2  # 2→4, 3→5, 4→6 (heading-4 always sub-section)
            if level > 6:
                level = 6
    document.add_heading(title, level=level)


def _b_list(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    style = _str(block.get("style")) or "bullet"
    items = block.get("items") or []
    if not items:
        return
    word_style = {
        "number": "List Number",
        "check": "List Bullet",
        "bullet": "List Bullet",
    }.get(style, "List Bullet")
    # Importer recognises list paragraphs by `<w:numPr>`; the built-in
    # 'List Bullet' / 'List Number' styles in python-docx's default
    # template don't include it, so we set numId/ilvl on every list
    # paragraph explicitly. numId values 1 (bullet) / 2 (number) are
    # arbitrary — round-trip only inspects existence + ilvl.
    num_id = "2" if style == "number" else "1"
    for item in items:
        if isinstance(item, dict):
            text = _str(item.get("text"))
            depth = int(item.get("depth") or 0)
        else:
            text = _str(item)
            depth = 0
        if not text and depth == 0:
            continue
        p = document.add_paragraph(style=word_style)
        _set_numpr(p, num_id=num_id, ilvl=str(depth))
        if depth > 0:
            p.paragraph_format.left_indent = Inches(0.25 * depth)
        prefix = ""
        if style == "check":
            prefix = "☐ "
        if prefix:
            run = p.add_run(prefix)
            run.font.size = Pt(11)
        _emit_inline_runs(p, text, ctx)


def _b_quote(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    text = _str(block.get("text"))
    cite = _str(block.get("cite"))
    if not (text or cite):
        return
    if text:
        # First line uses Intense Quote, subsequent lines use Quote.
        lines = text.splitlines() or [""]
        for i, line in enumerate(lines):
            style = "Intense Quote" if i == 0 else "Quote"
            try:
                p = document.add_paragraph(style=style)
            except KeyError:
                p = document.add_paragraph()
            _emit_inline_runs(p, line, ctx)
    if cite:
        try:
            p = document.add_paragraph(style="Quote")
        except KeyError:
            p = document.add_paragraph()
        run = p.add_run(f"— {cite}")
        run.italic = True


def _b_callout(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    variant = _str(block.get("variant")) or "info"
    title = _str(block.get("title"))
    text = _str(block.get("text"))
    if not (title or text):
        return
    bg, border = _CALLOUT_COLORS.get(variant, _CALLOUT_COLORS["info"])
    p = document.add_paragraph()
    _set_paragraph_shading(p, bg)
    _set_paragraph_left_border(p, border)
    if title:
        run = p.add_run(f"[{variant.upper()}] {title}")
        run.bold = True
    if text:
        if title:
            p.add_run("\n")
        for i, line in enumerate(text.splitlines() or [""]):
            if i > 0:
                p.add_run("\n")
            _emit_inline_runs_into_paragraph(p, line, ctx)


def _b_code(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    code = _str(block.get("code"))
    language = _str(block.get("language"))
    filename = _str(block.get("filename"))
    if filename or language:
        cap = document.add_paragraph()
        run = cap.add_run(filename or language or "code")
        run.italic = True
        run.font.size = Pt(9)
    if not code:
        return
    p = document.add_paragraph()
    _set_paragraph_shading(p, _CODE_SHADING)
    for i, line in enumerate(code.splitlines() or [""]):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def _b_math(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    expr = _str(block.get("expression"))
    if not expr:
        return
    display = _str(block.get("display")) or "block"
    delim = "$$" if display == "block" else "$"
    p = document.add_paragraph()
    run = p.add_run(f"{delim}{expr}{delim}")
    run.italic = True
    run.font.name = "Cambria Math"


def _b_table(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    cells = block.get("cells")
    if isinstance(cells, list) and cells:
        _emit_table_cells(document, block, cells, ctx)
    else:
        _emit_table_flat(document, block, ctx)
    # Caption (if meta.note set) — Word's "Caption" style follows the table.
    meta = block.get("meta") or {}
    caption = _str(meta.get("note")) if isinstance(meta, dict) else ""
    if caption and caption != "page-break-before":
        try:
            cap_p = document.add_paragraph(style="Caption")
        except KeyError:
            cap_p = document.add_paragraph()
        run = cap_p.add_run(caption)
        run.italic = True


def _emit_table_flat(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Legacy headers/rows layout — every slot is its own cell, no merges."""
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    if not headers and not rows:
        return
    if not headers:
        cols = max((len(r) for r in rows), default=0)
        headers = [""] * cols
    n_cols = len(headers)
    if n_cols == 0:
        return
    table = document.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for c, h in enumerate(headers):
        cell = hdr_cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(_str(h))
        run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        row_cells = table.rows[r_idx].cells
        for c_idx in range(n_cols):
            cell = row_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            value = _str(row[c_idx]) if c_idx < len(row) else ""
            _emit_inline_runs(p, value, ctx)


def _emit_table_cells(
    document: Any,
    block: dict[str, Any],
    cells: list[dict[str, Any]],
    ctx: _Ctx,
) -> None:
    """Merge-aware layout — uses python-docx's `Cell.merge` to mirror our
    sparse `cells` list back into Word ``<w:gridSpan>`` / ``<w:vMerge>`` so a
    DOCX → MX → DOCX round trip preserves the table shape.
    """
    if not cells:
        return
    max_r = 0
    max_c = 0
    for cell in cells:
        r_end = int(cell.get("r", 0)) + max(1, int(cell.get("rowSpan") or 1)) - 1
        c_end = int(cell.get("c", 0)) + max(1, int(cell.get("colSpan") or 1)) - 1
        max_r = max(max_r, r_end)
        max_c = max(max_c, c_end)
    n_rows = max_r + 1
    n_cols = max_c + 1
    if n_rows == 0 or n_cols == 0:
        return
    table = document.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # First fill every anchor cell with its text. Merges happen afterwards
    # because python-docx merges modify the underlying XML grid; touching
    # an already-merged cell is awkward, so we lay text down first.
    for cell in cells:
        r = int(cell.get("r", 0))
        c = int(cell.get("c", 0))
        text = _str(cell.get("text"))
        is_header = bool(cell.get("header"))
        try:
            tc = table.rows[r].cells[c]
        except IndexError:
            continue
        tc.text = ""
        p = tc.paragraphs[0]
        if is_header:
            run = p.add_run(text)
            run.bold = True
        else:
            _emit_inline_runs(p, text, ctx)

    # Now apply each anchor's merges. python-docx's Cell.merge takes a
    # diagonal-opposite cell and grows the rectangle; idempotent for 1×1.
    for cell in cells:
        r = int(cell.get("r", 0))
        c = int(cell.get("c", 0))
        rs = max(1, int(cell.get("rowSpan") or 1))
        cs = max(1, int(cell.get("colSpan") or 1))
        if rs == 1 and cs == 1:
            continue
        try:
            anchor = table.rows[r].cells[c]
            far = table.rows[r + rs - 1].cells[c + cs - 1]
            anchor.merge(far)
        except IndexError:
            continue


def _b_kpi_cards(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    items = block.get("items") or []
    if not items:
        return
    grid = items[:4]
    # 2x2 table with merged cells per KPI: label small, value big.
    n_rows = (len(grid) + 1) // 2
    table = document.add_table(rows=n_rows, cols=2)
    table.style = "Table Grid"
    for i, it in enumerate(grid):
        r = i // 2
        c = i % 2
        cell = table.rows[r].cells[c]
        cell.text = ""
        label_p = cell.paragraphs[0]
        run = label_p.add_run(_str(it.get("label")))
        run.font.size = Pt(10)
        value_p = cell.add_paragraph()
        run = value_p.add_run(_str(it.get("value")))
        run.bold = True
        run.font.size = Pt(18)
        delta = _str(it.get("delta"))
        trend = _str(it.get("trend"))
        if delta or trend:
            d_p = cell.add_paragraph()
            run = d_p.add_run(
                f"{delta} ({trend})" if delta and trend else delta or trend
            )
            run.font.size = Pt(9)


def _b_chart(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    title = _str(block.get("title"))
    chart_type = _str(block.get("chartType") or block.get("chart_type"))
    data = block.get("data") or {}
    labels = data.get("labels") or []
    series = data.get("series") or []
    p = document.add_paragraph()
    run = p.add_run(f"[차트] {title or chart_type or 'chart'}".strip())
    run.bold = True
    if not (labels or series):
        return
    # Render data as a small docx Table for fidelity.
    headers = ["계열", *[_str(label) for label in labels]]
    table = document.add_table(rows=1 + len(series), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for c, h in enumerate(headers):
        hdr[c].text = ""
        run = hdr[c].paragraphs[0].add_run(h)
        run.bold = True
    for r, sd in enumerate(series, start=1):
        row = table.rows[r].cells
        row[0].text = _str(sd.get("name"))
        values = sd.get("values") or []
        for c, v in enumerate(values, start=1):
            if c < len(headers):
                row[c].text = _str(v)
    note_p = document.add_paragraph()
    run = note_p.add_run("[차트 — DOCX export 는 데이터 표로만 보존]")
    run.italic = True
    run.font.size = Pt(9)


def _b_gantt(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    tasks = block.get("tasks") or []
    p = document.add_paragraph()
    run = p.add_run("[Gantt]")
    run.bold = True
    if not tasks:
        return
    for t in tasks:
        sub = document.add_paragraph(style="List Bullet")
        run = sub.add_run(
            f"{_str(t.get('name'))} : {_str(t.get('start'))} → {_str(t.get('end'))}"
        )
        run.font.size = Pt(11)


def _b_flow(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    engine = _str(block.get("engine")) or "mermaid"
    source = _str(block.get("source"))
    p = document.add_paragraph()
    run = p.add_run(f"[{engine} 다이어그램]")
    run.bold = True
    if source:
        code_p = document.add_paragraph()
        _set_paragraph_shading(code_p, _CODE_SHADING)
        for i, line in enumerate(source.splitlines() or [""]):
            if i > 0:
                code_p.add_run().add_break()
            r = code_p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(10)


def _b_org_chart(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    root = block.get("root") or {}
    if not root:
        return
    p = document.add_paragraph()
    run = p.add_run("[조직도]")
    run.bold = True

    def render_node(node: dict[str, Any], depth: int) -> None:
        label = _str(node.get("label"))
        role = _str(node.get("role"))
        line = f"{label} — {role}" if role else label
        if not line:
            return
        sub = document.add_paragraph(style="List Bullet")
        sub.paragraph_format.left_indent = Inches(0.25 * depth)
        sub.add_run(line)
        for child in node.get("children") or []:
            render_node(child, depth + 1)

    render_node(root, 0)


def _b_iframe(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    src = _str(block.get("src"))
    title = _str(block.get("title")) or src or "iframe"
    _emit_link_paragraph(document, "[iframe] ", title, src)


def _b_video(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    url = _str(block.get("url"))
    title = _str(block.get("title")) or url or "video"
    _emit_link_paragraph(document, "🎬 ", title, url)


def _b_image(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    image_id = _str(block.get("imageId") or block.get("image_id"))
    caption = _str(block.get("caption"))
    meta = block.get("meta") or {}
    width_px = None
    height_px = None
    if isinstance(meta, dict):
        try:
            if meta.get("width"):
                width_px = int(meta["width"])
        except (TypeError, ValueError):
            pass
        try:
            if meta.get("height"):
                height_px = int(meta["height"])
        except (TypeError, ValueError):
            pass

    resolver = ctx.opts.image_resolver
    info: dict[str, Any] | None = None
    if resolver and image_id:
        try:
            info = resolver(image_id)
        except Exception:
            info = None
    blob = info.get("bytes") if info else None
    if not blob:
        # Text fallback.
        p = document.add_paragraph()
        run = p.add_run(f"🖼 {caption or image_id or 'image'}")
        run.italic = True
        return
    p = document.add_paragraph()
    run = p.add_run()
    try:
        kwargs: dict[str, Any] = {}
        if width_px:
            kwargs["width"] = Inches(width_px / 96)  # 96 dpi
        if height_px and not width_px:
            kwargs["height"] = Inches(height_px / 96)
        run.add_picture(io.BytesIO(blob), **kwargs)
    except Exception:
        run.text = f"🖼 {caption or image_id or 'image'}"
        run.italic = True
        return
    if caption:
        try:
            cap_p = document.add_paragraph(style="Caption")
        except KeyError:
            cap_p = document.add_paragraph()
        ctx.figure_counter += 1
        prefix_run = cap_p.add_run(f"그림 {ctx.figure_counter}: ")
        prefix_run.bold = True
        run = cap_p.add_run(caption)
        run.italic = True

    # Image link — if present, append a small "↗ 링크 열기" hyperlink line
    # below the caption (python-docx can't wrap an inline-image hyperlink
    # without raw XML manipulation; this is the same compromise the FE
    # makes).
    link = _str(block.get("link"))
    if link:
        target = link if link.startswith(("http://", "https://")) else f"/docs/{link.lstrip('/')}"
        link_p = document.add_paragraph()
        _add_hyperlink(link_p, target, "↗ 링크 열기")


def _b_gallery(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    items = block.get("items") or []
    if not items:
        return
    for it in items:
        single = {
            "type": "image",
            "imageId": it.get("imageId") or it.get("image_id"),
            "caption": it.get("caption"),
            "alt": it.get("alt"),
        }
        _b_image(document, single, ctx)


def _b_file(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    file_id = _str(block.get("fileId") or block.get("file_id"))
    name = _str(block.get("name")) or file_id or "file"
    href = f"/api/v1/files/{file_id}/download" if file_id else ""
    _emit_link_paragraph(document, "📎 ", name, href)


def _b_doc_link_card(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    slug = _str(block.get("slug"))
    title = _str(block.get("title")) or slug
    p = document.add_paragraph()
    p.add_run("📄 ")
    if slug:
        _add_hyperlink(p, f"/docs/{slug}", title or slug)
    else:
        p.add_run(title)


def _b_glossary_ref(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    term = _str(block.get("term"))
    definition = _str(block.get("definition"))
    if not term:
        return
    p = document.add_paragraph()
    run = p.add_run(term)
    run.italic = True
    if definition:
        run2 = p.add_run(f" ({definition})")
        run2.font.size = Pt(10)


def _b_columns(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    columns = block.get("columns") or []
    for col in columns:
        for inner in col:
            _render_block(document, inner, ctx)


def _b_tabs(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    tabs = block.get("tabs") or []
    for tab in tabs:
        label = _str(tab.get("label"))
        if label:
            p = document.add_paragraph()
            run = p.add_run(f"▸ {label}")
            run.bold = True
        for inner in tab.get("blocks") or []:
            _render_block(document, inner, ctx)


def _b_accordion(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    items = block.get("items") or []
    for it in items:
        label = _str(it.get("label"))
        if label:
            p = document.add_paragraph()
            run = p.add_run(f"▾ {label}")
            run.bold = True
        for inner in it.get("blocks") or []:
            _render_block(document, inner, ctx)


def _b_data_source(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    endpoint = _str(block.get("endpoint"))
    render = _str(block.get("render"))
    p = document.add_paragraph()
    run = p.add_run(f"[데이터 소스] {endpoint} · render={render}")
    run.font.size = Pt(11)


def _b_dashboard_embed(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    provider = _str(block.get("provider"))
    panel = _str(block.get("panelId") or block.get("panel_id"))
    p = document.add_paragraph()
    run = p.add_run(f"[대시보드] {provider} · panel={panel}")
    run.font.size = Pt(11)


def _b_calculator(document: Any, block: dict[str, Any], _ctx: _Ctx) -> None:
    label = _str(block.get("label"))
    formula = _str(block.get("formula"))
    if not (label or formula):
        return
    p = document.add_paragraph()
    run = p.add_run(f"[계산기] {label}".strip())
    run.bold = True
    if formula:
        sub = document.add_paragraph()
        run = sub.add_run(f"수식: {formula}")
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def _b_bibliography(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Render a BibliographyBlock as a docx heading + numbered list. Word
    has no first-class "bibliography" type without a citation source XML;
    we settle for a styled "References" heading and one paragraph per
    entry so the output reads naturally and round-trips back via the
    Step-3 References heuristic on import.
    """
    title = _str(block.get("title")) or "참고문헌"
    entries = block.get("entries") or []
    if not isinstance(entries, list) or not entries:
        return
    try:
        h = document.add_paragraph(style="Heading 2")
    except KeyError:
        h = document.add_paragraph()
    h_run = h.add_run(title)
    h_run.bold = True
    for idx, entry in enumerate(entries, start=1):
        if not isinstance(entry, dict):
            continue
        text = _str(entry.get("text"))
        url = _str(entry.get("url"))
        key = _str(entry.get("key"))
        p = document.add_paragraph()
        # Number prefix + key (when set) + text + URL appended in italics.
        prefix = f"[{idx}] "
        if key:
            prefix = f"[{key}] "
        run = p.add_run(prefix)
        run.bold = True
        _emit_inline_runs(p, text, ctx)
        if url:
            p.add_run(" ")
            url_run = p.add_run(url)
            url_run.italic = True


# ── New block handlers (round-trip-safe text rendering for blocks
# that previously fell through to the unknown-block placeholder).
# Each one renders a human-readable representation that round-trips
# through DOCX import as best as possible.


def _b_pdf(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """PDF block — embed as a labeled link + page indicator (DOCX can't
    embed PDFs natively as inline objects via python-docx)."""
    file_id = _str(block.get("fileId") or block.get("file_id"))
    title = _str(block.get("title")) or file_id or "PDF"
    page = block.get("page")
    p = document.add_paragraph()
    run = p.add_run(f"📄 PDF: {title}")
    run.bold = True
    if page is not None:
        p.add_run(f" — page {page}")
    if file_id:
        p.add_run(f" (file_id: {file_id})").italic = True


def _b_whiteboard(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Whiteboard — list strokes / shapes / text labels as a bullet list.
    Real visual fidelity needs SVG embed which python-docx can't do
    natively; the text summary still preserves the data."""
    p = document.add_paragraph()
    p.add_run("🖍 Whiteboard").bold = True
    elems = block.get("elements") or []
    for el in elems:
        kind = el.get("type") or el.get("kind") or "?"
        label = el.get("label") or el.get("text") or kind
        coords = el.get("coords") or el.get("points") or ""
        item = document.add_paragraph(style="List Bullet")
        item.add_run(f"{kind}: {label}")
        if coords:
            item.add_run(f"  ({coords})").italic = True


def _b_form(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Form — title + each question on its own line with kind/required."""
    title = _str(block.get("title")) or "Form"
    description = _str(block.get("description"))
    h = document.add_paragraph()
    h.add_run(f"📝 {title}").bold = True
    if description:
        document.add_paragraph(description).italic = True
    for q in block.get("questions") or []:
        p = document.add_paragraph(style="List Number")
        label = _str(q.get("label")) or _str(q.get("id"))
        kind = _str(q.get("kind")) or "text"
        required = " *" if q.get("required") else ""
        p.add_run(f"{label}{required}").bold = True
        p.add_run(f"  [{kind}]").italic = True
        opts = q.get("options") or []
        if isinstance(opts, list) and opts:
            opts_p = document.add_paragraph()
            opts_p.add_run("    선택: " + ", ".join(str(o) for o in opts)).italic = True


def _b_quiz(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Quiz — title + each question with options and (admin only) correct answer."""
    title = _str(block.get("title")) or "Quiz"
    h = document.add_paragraph()
    h.add_run(f"❓ {title}").bold = True
    passing = block.get("passing_score")
    if passing is not None:
        document.add_paragraph(f"통과 점수: {passing}").italic = True
    for q in block.get("questions") or []:
        p = document.add_paragraph(style="List Number")
        label = _str(q.get("label")) or _str(q.get("id"))
        kind = _str(q.get("kind")) or "single"
        p.add_run(label).bold = True
        p.add_run(f"  [{kind}]").italic = True
        for opt in q.get("options") or []:
            document.add_paragraph(f"    • {opt}")
        # Show correct answer only when the requester is an editor/admin.
        if ctx.requester_role in ("editor", "admin", "owner"):
            correct = q.get("correct_answer") or q.get("correct")
            if correct is not None:
                cp = document.add_paragraph()
                cp.add_run(f"    정답: {correct}").italic = True
            expl = _str(q.get("explanation"))
            if expl:
                document.add_paragraph(f"    해설: {expl}").italic = True


def _b_image_annotation(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """ImageAnnotation — embed the base image + bullet-list annotations."""
    image_id = _str(block.get("imageId") or block.get("image_id"))
    p = document.add_paragraph()
    p.add_run("✏ Image annotation").bold = True
    # Try to embed the image if resolver gives us bytes.
    if image_id and ctx.options.image_resolver:
        info = ctx.options.image_resolver(image_id)
        if info and info.get("bytes"):
            try:
                from io import BytesIO
                document.add_picture(BytesIO(info["bytes"]))
            except Exception:
                pass
    for ann in block.get("annotations") or []:
        kind = ann.get("kind") or ann.get("type") or "marker"
        label = ann.get("label") or ""
        coords = ann.get("coords") or ann.get("box") or ann.get("point") or ""
        item = document.add_paragraph(style="List Bullet")
        item.add_run(f"{kind}: {label}")
        if coords:
            item.add_run(f"  ({coords})").italic = True


def _b_spacer(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Spacer — emit empty paragraphs to mimic the FE gap.
    sm=1 empty paragraph, md=2, lg=4."""
    size = _str(block.get("size")) or "md"
    count = {"sm": 1, "md": 2, "lg": 4}.get(size, 2)
    for _ in range(count):
        document.add_paragraph("")


def _b_spreadsheet(document: Any, block: dict[str, Any], ctx: _Ctx) -> None:
    """Spreadsheet — render as a native DOCX table. Formula cells show
    both the formula and (if computed) value via meta."""
    title = _str(block.get("title"))
    if title:
        h = document.add_paragraph()
        h.add_run(f"⊞ {title}").bold = True
    cells = block.get("cells") or {}
    headers = block.get("headers") or []
    rows = block.get("rows", 0)
    cols = block.get("cols", 0)
    if not rows or not cols:
        # derive from cell keys (e.g. "B2")
        max_col = max_row = 0
        for k in cells:
            if len(k) >= 2 and k[0].isalpha():
                c = ord(k[0].upper()) - ord("A") + 1
                try:
                    r = int(k[1:])
                except ValueError:
                    continue
                max_col = max(max_col, c)
                max_row = max(max_row, r)
        rows, cols = max_row, max_col
    if not rows or not cols:
        document.add_paragraph("(빈 스프레드시트)").italic = True
        return
    table = document.add_table(rows=rows + (1 if headers else 0), cols=cols)
    table.style = "Light Grid"
    r0 = 0
    if headers:
        for c, h in enumerate(headers[:cols]):
            table.rows[0].cells[c].text = str(h)
        r0 = 1
    for r in range(rows):
        for c in range(cols):
            key = f"{chr(ord('A') + c)}{r + 1}"
            v = cells.get(key)
            if isinstance(v, dict):
                # {value: ..., formula: ...} shape
                text_val = v.get("value", v.get("formula", ""))
                table.rows[r + r0].cells[c].text = str(text_val)
            elif v is not None:
                table.rows[r + r0].cells[c].text = str(v)


_BLOCK_HANDLERS: dict[str, Any] = {
    "paragraph": _b_paragraph,
    "heading-4": _b_heading_4,
    "list": _b_list,
    "quote": _b_quote,
    "callout": _b_callout,
    "code": _b_code,
    "math": _b_math,
    "table": _b_table,
    "kpi-cards": _b_kpi_cards,
    "chart": _b_chart,
    "gantt": _b_gantt,
    "flow": _b_flow,
    "org-chart": _b_org_chart,
    "iframe": _b_iframe,
    "video": _b_video,
    "image": _b_image,
    "gallery": _b_gallery,
    "file": _b_file,
    "doc-link-card": _b_doc_link_card,
    "glossary-ref": _b_glossary_ref,
    "columns": _b_columns,
    "tabs": _b_tabs,
    "accordion": _b_accordion,
    "data-source": _b_data_source,
    "dashboard-embed": _b_dashboard_embed,
    "calculator": _b_calculator,
    "bibliography": _b_bibliography,
    # ── newly added (Round 5) ──
    "pdf": _b_pdf,
    "whiteboard": _b_whiteboard,
    "form": _b_form,
    "quiz": _b_quiz,
    "image-annotation": _b_image_annotation,
    "spreadsheet": _b_spreadsheet,
    # ── newly added (Round 7) ──
    "spacer": _b_spacer,
}


# ── Inline run emitter ───────────────────────────────────────────────


def _emit_inline_runs(paragraph: Any, text: str, ctx: _Ctx) -> None:
    """Tokenize markdown-lite and emit corresponding docx runs.

    Handles **bold**, *italic*, ~~strike~~, __underline__, `code`, [label](url),
    [[slug]] / [[slug|label]], [^N] footnote markers.
    """
    _emit_inline_runs_into_paragraph(paragraph, text, ctx)


def _emit_inline_runs_into_paragraph(paragraph: Any, text: str, ctx: _Ctx) -> None:
    if not text:
        return
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        groups = m.groupdict()
        if groups.get("bold"):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif groups.get("strike"):
            run = paragraph.add_run(m.group(4))
            run.font.strike = True
        elif groups.get("underline"):
            run = paragraph.add_run(m.group(6))
            run.underline = True
        elif groups.get("italic"):
            run = paragraph.add_run(m.group(8))
            run.italic = True
        elif groups.get("code"):
            run = paragraph.add_run(m.group(10))
            run.font.name = "Consolas"
            _set_run_shading(run, _INLINE_CODE_SHADING)
        elif groups.get("link"):
            label = m.group(12)
            url = m.group(13)
            _add_hyperlink(paragraph, url, label)
        elif groups.get("wiki"):
            inner = m.group(14)
            if "|" in inner:
                slug, _, label = inner.partition("|")
            else:
                slug = inner
                label = inner
            slug = slug.strip()
            label = label.strip() or slug
            href = slug if slug.startswith("#") else f"/docs/{slug}"
            _add_hyperlink(paragraph, href, label)
        elif groups.get("footnote"):
            marker = m.group(16)
            body = ctx.footnotes.get(marker) or marker
            _add_footnote(paragraph, body)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ── XML helpers ──────────────────────────────────────────────────────


def _set_paragraph_shading(paragraph: Any, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_paragraph_left_border(paragraph: Any, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")  # 1/8 pt units → 3pt
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_numpr(paragraph: Any, *, num_id: str, ilvl: str) -> None:
    """Add `<w:numPr>` to the paragraph so importers see it as a list item."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    numPr.append(ilvl_el)
    numPr.append(num_id_el)
    pPr.append(numPr)


def _set_run_shading(run: Any, hex_color: str) -> None:
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _add_hyperlink(paragraph: Any, url: str, label: str) -> Any:
    """Insert a clickable hyperlink (preserves URL in the .docx).

    python-docx 가 직접 노출하지 않으므로 part.relate_to + raw XML 으로 구성.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F3AA8")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    text = OxmlElement("w:t")
    text.text = label
    text.set(qn("xml:space"), "preserve")
    new_run.append(text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_footnote(paragraph: Any, body: str) -> None:
    """Append a marker run with the footnote body in parentheses.

    python-docx 가 footnote part 직접 추가를 노출하지 않아 round-trip 성을
    보장하는 가장 안전한 방식은 'inline 표기' (markdown-lite 와 동일). 본문은
    `[^marker]` 형식 그대로 둔다 — importer 가 다시 인식 가능.
    """
    run = paragraph.add_run(f" ({body})")
    run.font.size = Pt(9)
    run.font.superscript = True


def _emit_link_paragraph(
    document: Any, prefix: str, label: str, url: str
) -> None:
    p = document.add_paragraph()
    if prefix:
        p.add_run(prefix)
    if url:
        _add_hyperlink(p, url, label or url)
    else:
        p.add_run(label)


# ── Footnote collection ─────────────────────────────────────────────


_FOOTNOTE_DEF_RE = re.compile(r"^\s*\[\^([^\]]+)\]:\s*(.*)$")


def _is_footnote_def(text: str) -> bool:
    return bool(_FOOTNOTE_DEF_RE.match(text or ""))


def _collect_footnotes(doc: dict[str, Any]) -> dict[str, str]:
    """Walk doc and gather `[^N]: body` definitions from paragraph blocks."""
    out: dict[str, str] = {}

    def walk_blocks(blocks: list[dict[str, Any]]) -> None:
        for b in blocks or []:
            if b.get("type") == "paragraph":
                m = _FOOTNOTE_DEF_RE.match(_str(b.get("text")))
                if m:
                    out[m.group(1)] = m.group(2).strip()
            for col in b.get("columns") or []:
                walk_blocks(col)
            for tab in b.get("tabs") or []:
                walk_blocks(tab.get("blocks") or [])
            for it in b.get("items") or []:
                if isinstance(it, dict) and it.get("blocks"):
                    walk_blocks(it["blocks"])

    def walk_sections(secs: list[dict[str, Any]]) -> None:
        for s in secs or []:
            walk_blocks(s.get("blocks") or [])
            walk_sections(s.get("subsections") or [])

    walk_sections(doc.get("sections") or [])
    return out


# ── Helpers ──────────────────────────────────────────────────────────


def _str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, bool):
        return "true" if v else "false"
    return str(v)
