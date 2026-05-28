"""DOCX export 렌더러 단위 테스트.

renderer 자체는 부수효과 없는 pure function. endpoint 통합은 한 건만 스모크.
.docx 는 zip 컨테이너이므로:
  - magic bytes (PK\\x03\\x04) 로 zip 여부 검증
  - python-docx 로 다시 열어 단락/표/스타일 확인
"""
from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from docx.document import Document as DocxDocument
from httpx import ASGITransport, AsyncClient

from app.main import app
from app.services.docx_export import DocxOptions, render_docx


def _doc(blocks: list[dict] | None = None, **overrides: object) -> dict:
    base = {
        "schema_version": "1.0",
        "id": "01TESTDOC0000000000000000Z",
        "slug": "fixture-docx",
        "title": "Word 픽스처",
        "summary": "단위 테스트 픽스처입니다.",
        "metadata": {
            "division": "MX",
            "team": "Editor",
            "owners": ["someone@example.com"],
            "tags": ["unit", "docx"],
            "confidentiality": "internal",
        },
        "sections": [
            {
                "id": "01SEC00000000000000000000A",
                "number": "1",
                "level": 1,
                "title": "서론",
                "blocks": blocks or [
                    {
                        "type": "paragraph",
                        "id": "01P000000000000000000000A1",
                        "text": "첫 문단입니다.",
                    }
                ],
                "subsections": [],
            }
        ],
    }
    base.update(overrides)
    return base


def _all_text(doc: DocxDocument) -> str:
    """문서 모든 paragraph + table 셀 텍스트를 모아 검색용으로 반환."""
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ── pure renderer ────────────────────────────────────────────────────


def test_renderer_returns_valid_docx_zip() -> None:
    out = render_docx(_doc())
    assert out[:4] == b"PK\x03\x04"
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        names = zf.namelist()
        assert "word/document.xml" in names


def test_renderer_emits_title_and_summary() -> None:
    out = render_docx(_doc())
    doc = Document(io.BytesIO(out))
    text = _all_text(doc)
    assert "Word 픽스처" in text
    assert "단위 테스트 픽스처입니다." in text
    assert "MX" in text  # division


def test_renderer_emits_section_headings() -> None:
    out = render_docx(_doc())
    doc = Document(io.BytesIO(out))
    headings = [
        p.text for p in doc.paragraphs
        if p.style is not None
        and p.style.name is not None
        and p.style.name.startswith("Heading")
    ]
    assert any("1 서론" in h for h in headings)


def test_renderer_paragraph_inline_bold_italic_strike() -> None:
    blocks = [
        {
            "type": "paragraph",
            "id": "01P00000000000000000010A",
            "text": "이것은 **굵게** 와 *기울임* 그리고 ~~취소~~ 입니다.",
        }
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    bold_runs = []
    italic_runs = []
    strike_runs = []
    for p in doc.paragraphs:
        for run in p.runs:
            if run.bold:
                bold_runs.append(run.text)
            if run.italic:
                italic_runs.append(run.text)
            if run.font.strike:
                strike_runs.append(run.text)
    assert "굵게" in bold_runs
    assert "기울임" in italic_runs
    assert "취소" in strike_runs


def test_renderer_table_emits_native_table_with_bold_header() -> None:
    blocks = [
        {
            "type": "table",
            "id": "01T1A0000000000000000001A",
            "headers": ["이름", "수량"],
            "rows": [["사과", "3"], ["배", "2"]],
            "meta": {"note": "재고표"},
        }
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert t.cell(0, 0).text == "이름"
    assert t.cell(1, 1).text == "3"
    # First-row runs should be bold.
    hdr = t.rows[0].cells[0].paragraphs[0]
    assert any(r.bold for r in hdr.runs)
    # Caption paragraph after the table.
    text = _all_text(doc)
    assert "재고표" in text


def test_renderer_list_styles_bullet_and_number() -> None:
    blocks = [
        {
            "type": "list",
            "id": "01L0000000000000000000001",
            "style": "number",
            "items": ["하나", "둘"],
        },
        {
            "type": "list",
            "id": "01L0000000000000000000002",
            "style": "check",
            "items": ["TODO1"],
        },
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    styles = [p.style.name for p in doc.paragraphs if p.style is not None]
    assert "List Number" in styles
    assert "List Bullet" in styles
    text = _all_text(doc)
    assert "하나" in text and "둘" in text
    assert "☐" in text and "TODO1" in text


def test_renderer_code_block_uses_monospace_font() -> None:
    blocks = [
        {
            "type": "code",
            "id": "01CD000000000000000000001",
            "language": "python",
            "code": "print('hello')",
        }
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    found = False
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name == "Consolas" and "hello" in run.text:
                found = True
    assert found


def test_renderer_quote_uses_quote_style() -> None:
    blocks = [
        {
            "type": "quote",
            "id": "01Q1000000000000000000001",
            "text": "측정할 수 없으면 관리할 수 없다.",
            "cite": "Drucker",
        }
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    styles = [
        p.style.name for p in doc.paragraphs
        if p.style is not None and p.style.name is not None
    ]
    assert any("Quote" in s for s in styles)
    text = _all_text(doc)
    assert "Drucker" in text


def test_renderer_callout_emits_paragraph_shading() -> None:
    blocks = [
        {
            "type": "callout",
            "id": "01CA000000000000000000001",
            "variant": "warn",
            "title": "주의",
            "text": "확인하세요.",
        }
    ]
    out = render_docx(_doc(blocks))
    # Inspect the raw document.xml for the shading element.
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert "w:shd" in doc_xml
    assert "FFF7E0" in doc_xml.upper() or "fff7e0" in doc_xml


def _make_png(size: int = 4) -> bytes:
    """Build a tiny but valid PNG via Pillow (already a project dep)."""
    from PIL import Image as _PILImage

    img = _PILImage.new("RGB", (size, size), (255, 0, 0))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def test_renderer_image_uses_resolver_bytes_when_available() -> None:
    png = _make_png()

    def resolver(image_id: str) -> dict | None:
        if image_id == "01IMG00000000000000000000":
            return {"bytes": png, "mime": "image/png"}
        return None

    blocks = [
        {
            "type": "image",
            "id": "01I00000000000000000001A",
            "imageId": "01IMG00000000000000000000",
            "caption": "그림 1",
        }
    ]
    out = render_docx(_doc(blocks), options=DocxOptions(image_resolver=resolver))
    doc = Document(io.BytesIO(out))
    # Image lives inside an inline shape — count via doc.inline_shapes.
    assert len(doc.inline_shapes) >= 1
    assert "그림 1" in _all_text(doc)


def test_renderer_image_without_resolver_falls_back_to_text() -> None:
    blocks = [
        {
            "type": "image",
            "id": "01I00000000000000000002A",
            "imageId": "01IMG00000000000000000999",
            "caption": "그림 2",
        }
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    text = _all_text(doc)
    assert "그림 2" in text
    assert len(doc.inline_shapes) == 0


def test_renderer_page_break_inserts_page_break_run() -> None:
    blocks = [
        {
            "type": "paragraph",
            "id": "01P00000000000000000050A",
            "text": "",
            "meta": {"note": "page-break-before"},
        },
        {
            "type": "paragraph",
            "id": "01P00000000000000000051A",
            "text": "다음 페이지",
        },
    ]
    out = render_docx(_doc(blocks))
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert 'w:type="page"' in doc_xml


def test_renderer_speaker_note_dropped() -> None:
    blocks = [
        {
            "type": "paragraph",
            "id": "01P00000000000000000060A",
            "text": "본문",
            "meta": {"note": "speaker: 강조"},
        }
    ]
    out = render_docx(_doc(blocks))
    doc = Document(io.BytesIO(out))
    # speaker note paragraph should be dropped entirely.
    assert "본문" not in _all_text(doc)


def test_renderer_hyperlink_link_is_external_relationship() -> None:
    blocks = [
        {
            "type": "paragraph",
            "id": "01P00000000000000000070A",
            "text": "[홈페이지](https://example.com) 를 보세요.",
        }
    ]
    out = render_docx(_doc(blocks))
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "https://example.com" in rels
    assert "hyperlink" in rels.lower()


def test_renderer_table_stripe_false_uses_plain_style() -> None:
    """`options.stripe=False` → Table Grid (plain); True (default) → Light Grid Accent 1."""
    striped_blocks = [
        {
            "type": "table",
            "id": "01TBL00000000000000000001",
            "headers": ["A", "B"],
            "rows": [["1", "2"]],
            "options": {"stripe": True},
        }
    ]
    plain_blocks = [
        {
            "type": "table",
            "id": "01TBL00000000000000000002",
            "headers": ["A", "B"],
            "rows": [["1", "2"]],
            "options": {"stripe": False},
        }
    ]
    striped_out = render_docx(_doc(striped_blocks))
    plain_out = render_docx(_doc(plain_blocks))
    striped_doc = Document(io.BytesIO(striped_out))
    plain_doc = Document(io.BytesIO(plain_out))
    striped_style = striped_doc.tables[0].style.name if striped_doc.tables[0].style else ""
    plain_style = plain_doc.tables[0].style.name if plain_doc.tables[0].style else ""
    assert "Light Grid" in striped_style
    assert plain_style == "Table Grid"


def test_renderer_callout_emits_hidden_marker_run() -> None:
    """CalloutBlock → hidden `Widget: callout (variant)` marker run for round-trip."""
    blocks = [
        {
            "type": "callout",
            "id": "01CALLOUT0000000000000001",
            "variant": "warn",
            "title": "주의",
            "text": "확인하세요.",
        }
    ]
    out = render_docx(_doc(blocks))
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert "Widget: callout" in doc_xml
    # python-docx writes hidden runs with `<w:vanish/>`.
    assert "vanish" in doc_xml


def test_renderer_image_width_enum_drives_picture_size() -> None:
    """`width=lg` → ~6.25 inch wide Picture (600 px / 96 dpi); `full` leaves intrinsic."""
    png = _make_png(size=8)

    def resolver(image_id: str) -> dict | None:
        return {"bytes": png, "mime": "image/png"}

    lg_blocks = [
        {
            "type": "image",
            "id": "01I00000000000000000010A",
            "imageId": "01IMG00000000000000000100",
            "caption": "lg",
            "width": "lg",
        }
    ]
    full_blocks = [
        {
            "type": "image",
            "id": "01I00000000000000000011A",
            "imageId": "01IMG00000000000000000101",
            "caption": "full",
            "width": "full",
        }
    ]
    lg_out = render_docx(_doc(lg_blocks), options=DocxOptions(image_resolver=resolver))
    full_out = render_docx(_doc(full_blocks), options=DocxOptions(image_resolver=resolver))
    lg_doc = Document(io.BytesIO(lg_out))
    full_doc = Document(io.BytesIO(full_out))
    # python-docx exposes shape width in EMU (1 inch = 914400 EMU).
    lg_w = lg_doc.inline_shapes[0].width
    full_w = full_doc.inline_shapes[0].width
    # lg ≈ 600 / 96 = 6.25 inches → 5_715_000 EMU (allow ±5% rounding).
    assert 5_400_000 <= lg_w <= 6_000_000
    # full leaves intrinsic resolution — for an 8px PNG that's tiny (~75_000 EMU),
    # so anything below 1 inch confirms we did NOT scale to 6.25 in.
    assert full_w < 914_400


# ── endpoint integration ─────────────────────────────────────────────


SEED_SLUG = "month-end-closing"


@pytest.mark.asyncio
async def test_export_docx_endpoint_returns_attachment() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post("/api/v1/exports/docx", json={"slug": SEED_SLUG})
    assert r.status_code == 200, r.text
    ctype = r.headers.get("content-type", "")
    assert "wordprocessingml.document" in ctype
    cd = r.headers.get("content-disposition") or ""
    assert "attachment" in cd
    assert f"{SEED_SLUG}.docx" in cd
    body = r.content
    assert body[:4] == b"PK\x03\x04"


@pytest.mark.asyncio
async def test_export_docx_endpoint_404_for_missing_slug() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post(
            "/api/v1/exports/docx", json={"slug": "no-such-slug-xxx-xxx"}
        )
    assert r.status_code == 404


@pytest.mark.asyncio
async def test_export_docx_endpoint_rejects_missing_slug() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post("/api/v1/exports/docx", json={})
    assert r.status_code == 422


# ── {{var}} substitution parity with markdown / html / pptx ─────────


def test_renderer_substitutes_variable_tokens_in_paragraphs() -> None:
    """docx export should resolve `{{var}}` tokens just like the other
    three renderers. Without this the docx leg of round-trip emitted
    literal `{{name}}` tokens — a regression caught only by manual
    inspection.
    """
    doc = _doc(
        blocks=[
            {
                "type": "paragraph",
                "id": "01P000000000000000000000A1",
                "text": "안녕하세요 {{user}}",
            }
        ],
        variables={"user": "Park"},
    )
    out = render_docx(doc)
    text = _all_text(Document(io.BytesIO(out)))
    assert "안녕하세요 Park" in text
    assert "{{user}}" not in text


def test_renderer_substitutes_section_titles() -> None:
    doc = _doc(variables={"section_title": "본문"})
    doc["sections"][0]["title"] = "{{section_title|기본}}"
    out = render_docx(doc)
    text = _all_text(Document(io.BytesIO(out)))
    assert "본문" in text
    assert "{{section_title" not in text


def test_renderer_substitutes_fallback_when_variable_missing() -> None:
    doc = _doc(
        blocks=[
            {
                "type": "paragraph",
                "id": "01P000000000000000000000A2",
                "text": "오늘은 {{date|TBD}}.",
            }
        ],
        variables={},
    )
    out = render_docx(doc)
    text = _all_text(Document(io.BytesIO(out)))
    assert "오늘은 TBD." in text


def test_renderer_leaves_unfilled_tokens_literal() -> None:
    """Unfilled tokens with no fallback are preserved as `{{name}}` —
    matches the helper's contract so the user sees what's missing.
    """
    doc = _doc(
        blocks=[
            {
                "type": "paragraph",
                "id": "01P000000000000000000000A3",
                "text": "값: {{missing}}",
            }
        ],
        variables={},
    )
    out = render_docx(doc)
    text = _all_text(Document(io.BytesIO(out)))
    assert "{{missing}}" in text


def test_renderer_skips_substitution_inside_code_blocks() -> None:
    doc = _doc(
        blocks=[
            {
                "type": "code",
                "id": "01C000000000000000000000A4",
                "language": "python",
                "code": "secret = {{secret}}",
            }
        ],
        variables={"secret": "topsecret"},
    )
    out = render_docx(doc)
    text = _all_text(Document(io.BytesIO(out)))
    assert "{{secret}}" in text
    assert "topsecret" not in text


# ── TOC (include_toc opt-in) ─────────────────────────────────────────


def _multi_section_doc() -> dict:
    """두 개의 top-level section + 한 개의 subsection — TOC 라벨이 모두
    뽑히는지 검증할 픽스처."""
    return {
        "schema_version": "1.0",
        "id": "01TESTTOC0000000000000000Z",
        "slug": "fixture-toc",
        "title": "TOC 테스트",
        "summary": "",
        "metadata": {
            "division": "MX",
            "owners": ["someone@example.com"],
            "confidentiality": "internal",
        },
        "sections": [
            {
                "id": "01SECTOC00000000000000001",
                "number": "1",
                "level": 1,
                "title": "서론",
                "blocks": [],
                "subsections": [
                    {
                        "id": "01SECTOC00000000000000002",
                        "number": "1.1",
                        "level": 2,
                        "title": "배경",
                        "blocks": [],
                        "subsections": [],
                    }
                ],
            },
            {
                "id": "01SECTOC00000000000000003",
                "number": "2",
                "level": 1,
                "title": "결론",
                "blocks": [],
                "subsections": [],
            },
        ],
    }


def test_renderer_include_toc_default_off_no_toc_heading() -> None:
    """``include_toc`` 디폴트 False — '목차' 헤딩 안 박힘 (호환성)."""
    out = render_docx(_multi_section_doc())
    text = _all_text(Document(io.BytesIO(out)))
    assert "목차" not in text


def test_renderer_include_toc_opt_in_emits_section_titles() -> None:
    """``include_toc=True`` 면 title 직후에 '목차' 헤딩 + level 1/2 단락이
    나오고, 본문 헤딩은 그 뒤에 정상 emit."""
    out = render_docx(
        _multi_section_doc(), options=DocxOptions(include_toc=True)
    )
    doc = Document(io.BytesIO(out))
    text = _all_text(doc)
    # TOC heading + entries
    assert "목차" in text
    assert "1 서론" in text
    assert "1.1 배경" in text
    assert "2 결론" in text
    # 본문 헤딩이 여전히 emit
    headings = [
        p.text for p in doc.paragraphs
        if p.style is not None
        and p.style.name is not None
        and p.style.name.startswith("Heading")
    ]
    # '목차' 가 첫 헤딩, 그 뒤에 본문 섹션 헤딩
    assert headings[0] == "목차"
    assert any("1 서론" in h for h in headings[1:])
    assert any("2 결론" in h for h in headings[1:])
